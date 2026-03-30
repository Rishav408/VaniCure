from docx import Document
from docx.shared import Pt
import os

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    return doc.add_paragraph(text)

def create_documentation():
    doc = Document()

    # Title
    t = doc.add_heading('VaniCure — AI Respiratory Diagnostics Platform', 0)
    
    # Introduction
    add_paragraph(doc, "VaniCure is an advanced offline-first medical screening application. "
                       "It utilizes modern deep learning to analyze patient respiratory audio "
                       "(e.g., coughs, breathing) and predict potential risks for Tuberculosis (TB) "
                       "and Asthma. Designed for edge nodes in remote clinics, it runs full inference "
                       "locally on the device (zero cloud dependencies), protecting patient privacy.")
                       
    # Core Features
    add_heading(doc, 'Core Features', 1)
    
    doc.add_paragraph("Multi-Model Diagnostic Pipeline", style='List Bullet')
    p1 = doc.add_paragraph(style='List Bullet 2')
    p1.add_run("PANNs (CNN14): ").bold = True
    p1.add_run("Large-scale robust feature extractor trained on AudioSet, repurposed for deep respiratory screening. (527 classes, 312MB)")
    
    p2 = doc.add_paragraph(style='List Bullet 2')
    p2.add_run("YAMNet: ").bold = True
    p2.add_run("Ultra-fast MobileNet-based lightweight classifier for immediate anomaly recognition. (521 classes, 3.7MB)")
    
    p3 = doc.add_paragraph(style='List Bullet 2')
    p3.add_run("Custom CNN-BiLSTM: ").bold = True
    p3.add_run("Purpose-built architecture featuring a spatial CNN encoder and temporal BiLSTM sequence evaluator. Achieves >0.87 F1 on targeted respiratory sets.")

    p4 = doc.add_paragraph("Offline First & Secure: ", style='List Bullet')
    p4.add_run("FastAPI backend loads models locally sending zero data to the cloud.")
    
    p5 = doc.add_paragraph("Dynamic Outbreak Alerts: ", style='List Bullet')
    p5.add_run("Anomaly detection algorithm reads internal local storage looking for spatial clusters (≥2 critical cases within a parameter window).")
    
    p6 = doc.add_paragraph("Patient Records Pipeline: ", style='List Bullet')
    p6.add_run("React-driven dashboard stores patient meta-data alongside raw AI risk-scores locally, exportable securely to standard .csv files.")

    # Tech Stack
    add_heading(doc, 'Tech Stack', 1)
    doc.add_paragraph("Frontend: React 18, Vite, TypeScript, Tailwind CSS, Lucide Icons, Recharts (local state via localStorage)", style='List Bullet')
    doc.add_paragraph("Backend: Python 3.11, FastAPI, Uvicorn", style='List Bullet')
    doc.add_paragraph("AI / Machine Learning: PyTorch, TensorFlow, librosa, NumPy", style='List Bullet')

    # Architecture
    add_heading(doc, 'System Architecture — End-to-End Workflow', 1)
    add_paragraph(doc, "1. Patient Registration & Audio Input:")
    p = doc.add_paragraph()
    p.add_run("The patient arrives at a rural clinic. A health worker uses the React Frontend (localhost:3000) to register the patient's metadata and captures a live audio recording of their cough via microphone or file upload (.wav, .mp3).")
    
    add_paragraph(doc, "2. Audio Processing Pipeline:")
    p = doc.add_paragraph()
    p.add_run("The FastAPI backend receives the audio Blob. Before inference, the raw waveform is resampled to the model's required sample rate (e.g., 32kHz for PANNs, 16kHz for YAMNet, 22kHz for CNN-BiLSTM). Log-Mel Spectrograms and MFCC features are extracted and normalized.")
    
    add_paragraph(doc, "3. AI Model Pipeline (Sequential Execution):")
    p = doc.add_paragraph()
    p.add_run("The features are fed sequentially into PANNs CNN14, YAMNet, and the Custom CNN-BiLSTM. Each model outputs classification probabilities across numerous classes (cough, wheeze, breathing, throat clear).")
    
    add_paragraph(doc, "4. Risk Score Calculation:")
    p = doc.add_paragraph()
    p.add_run("The raw probabilities are combined securely. For example:\n")
    p.add_run("- TB Risk = cough×0.55 + throat×0.20 + wheeze×0.15 + breathing×0.10\n")
    p.add_run("- Asthma Risk = wheeze×0.50 + breathing×0.30 + cough×0.10\n")
    p.add_run("- Normal = 1.0 - (TB + Asthma)")

    add_paragraph(doc, "5. Multilingual Triage Agent:")
    p = doc.add_paragraph()
    p.add_run("A rule-based or SLM-powered agent conducts 2-3 rounds of follow up questions intelligently in Hindi, Gujarati, Marathi, or English based on patient preference.")
    
    add_paragraph(doc, "6. Final Output & Storage:")
    p = doc.add_paragraph()
    p.add_run("An ensemble average determines patient status:\n")
    p.add_run("- HIGH RISK (TB > 25%): Immediate hospital referral.\n")
    p.add_run("- ASTHMA (Asthma > 30%): Pulmonologist within 48 hours.\n")
    p.add_run("- LOW RISK (Normal): Routine checkup in 3 months.\n")
    p.add_run("Records are written seamlessly to offline local storage, where rapid spatial outbreak detection takes place.")

    # Save document
    save_path = '/Users/kunaltailor/Desktop/VaniCure_Project_Documentation.docx'
    doc.save(save_path)
    print(f"Document saved to {save_path}")

if __name__ == '__main__':
    create_documentation()
